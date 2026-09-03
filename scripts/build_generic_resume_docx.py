"""Build a general-purpose resume for Java backend, technical leadership and AI engineering roles."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"D:/github_dgy/output/docx/狄广宇-通用研发技术负责人版.docx")
FONT = "Microsoft YaHei"
NAVY = RGBColor(20, 55, 89)
BLUE = RGBColor(35, 103, 153)
INK = RGBColor(40, 44, 48)
MUTED = RGBColor(102, 109, 117)
PALE = "EAF2F8"
RULE = "B8CBD9"


def font(run, size=9.7, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def p_border(paragraph, color=RULE, size="8"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title.upper())
    font(r, 12.5, True, NAVY)
    p_border(p)
    return p


def add_bullet(doc, text, size=9.5, after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.52)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.keep_together = True
    if "：" in text and text.index("：") <= 9:
        head, tail = text.split("：", 1)
        r = p.add_run(head + "：")
        font(r, size, True, BLUE)
        r = p.add_run(tail)
        font(r, size)
    else:
        r = p.add_run(text)
        font(r, size)
    return p


def add_label_line(doc, label, text, size=9.6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.keep_together = True
    r = p.add_run(label)
    font(r, size, True, BLUE)
    r = p.add_run(text)
    font(r, size)
    return p


def add_project(doc, title, company, time, role, objective, bullets, stack=None, compact=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if not compact else 4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    font(r, 11.1 if not compact else 10.2, True, NAVY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.keep_with_next = True
    r = p2.add_run(company + "  |  " + time + "  |  " + role)
    font(r, 8.6 if compact else 8.8, False, MUTED)
    add_label_line(doc, "目标与范围：", objective, 9.35 if not compact else 8.9)
    for b in bullets:
        add_bullet(doc, b, 8.9 if compact else 9.5, after=0 if compact else 2)
    if stack:
        add_label_line(doc, "技术栈：", stack, 8.55 if compact else 9.1)


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.25)
    sec.bottom_margin = Cm(1.15)
    sec.left_margin = Cm(1.55)
    sec.right_margin = Cm(1.55)
    sec.header_distance = Cm(0.65)
    sec.footer_distance = Cm(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08
    for style_name in ("List Bullet", "List Bullet 2"):
        st = doc.styles[style_name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(9.5)

    # Quiet running furniture.
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("狄广宇  ·  Java 后端 / 技术负责人 / AI 应用工程化")
    font(hr, 8, False, MUTED)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("简历｜通用研发方向")
    font(fr, 8, False, MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("狄广宇")
    font(r, 24, True, NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Java 后端 · 技术负责人 · To B 平台 · AI 应用工程化")
    font(r, 13, False, BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("电话：19290569945　邮箱：2537676793@qq.com　GitHub：github.com/dgy-github")
    font(r, 9.5, False, MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("12 年研发经验　·　9 年 Java 后端及架构经验　·　本科（软件工程）")
    font(r, 10.3, True, INK)

    add_section(doc, "职业定位")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run("以 Java 后端与企业级系统交付为主线，能够把复杂业务拆成可配置、可扩展、可观测的产品能力；近年将 RAG、Agent、Workflow、工具调用和评测机制接入真实业务，适合技术负责人、Java 后端负责人、To B 平台研发及 AI 应用工程化岗位。")
    font(r, 10.2, False, INK)
    add_bullet(doc, "交付闭环：从需求沟通、方案设计、数据模型、接口与任务拆解，到开发、测试、上线、监控和持续迭代。")
    add_bullet(doc, "工程底座：Spring Boot、MyBatis、Spring Cloud、MySQL/Redis，覆盖微服务、数据同步、权限、幂等、消息和性能优化。")
    add_bullet(doc, "AI 增强：制造业排产 Agent、面料多模态检索、Coding Agent 和多 Agent 治理框架，强调安全边界、可回放、可评测和可验收。")
    add_bullet(doc, "团队协作：带领 2-3 人小组（1 名实习生、1 名前端），负责评审、拆解、Review、进度推进和交付验收；长期使用 Docker、Jenkins。")

    add_section(doc, "核心能力")
    add_label_line(doc, "后端与架构：", "Java 8/11/21、JVM、多线程、IO/NIO、Spring Boot、Spring Cloud、MyBatis、微服务、事件驱动、幂等与分布式一致性。")
    add_label_line(doc, "业务与平台：", "SRM、ERP、CRM、B2B 电商、支付、供应链、物联网、权限与审批；擅长规则抽象、数据建模、第三方集成和性能治理。")
    add_label_line(doc, "AI 应用工程化：", "LangGraph、RAG、Hybrid/GraphRAG、Tool Calling、MCP、Workflow、Prompt、Qwen、多模态检索、Langfuse、Golden Cases。")
    add_label_line(doc, "交付与质量：", "需求拆解、代码 Review、测试设计、E2E、日志追踪、权限审计、失败重试、回滚边界、Docker、Jenkins。")

    add_section(doc, "工作经历")
    add_label_line(doc, "浙江思维特数字科技有限公司　|　AI 应用研发 / 技术经理　|　2023.07-2026.04", "负责 SRM、ERP、CRM 产品线及 AI 应用方案、架构和团队交付。")
    add_label_line(doc, "工路（杭州）信息技术有限公司　|　Java 研发工程师 / 技术经理　|　2021.06-2023.07", "负责支付微服务、工业物联网平台和材料商城 B2B 电商的核心开发与上线。")
    add_label_line(doc, "杭州亿凯软件有限公司等　|　Java 软件工程师 / 模块负责人　|　2015.07-2021.06", "参与金融交易、银行智能投顾和政务 OA，承担后端模块、批处理、性能优化和现场交付。")
    add_label_line(doc, "教育背景：", "苏州工学院｜软件工程｜本科（统招）｜2010.09-2014.06", 9.2)

    page_break(doc)
    add_section(doc, "重点项目")
    add_project(doc, "制造业排产 Agent", "浙江思维特数字科技有限公司", "2025-2026.04", "技术负责人｜带领 2-3 人小组", "将生产订单查询、机台分析、规格推荐和排产诊断接入 Agent 工作流，保持核心业务动作可控、可追溯。", [
        "数据与路由：基于 49,981 条真实订单抽象 20 类意图；规则与 SQL 优先，约 80% 请求无需调用 LLM，离线意图识别准确率约 95%。",
        "业务建模：覆盖 75D / 100D / 300D 等规格排产、机台负载、替代机台和排产原因解释，把复杂规则固化为可复用服务能力。",
        "Agent 编排：用 LangGraph StateGraph 编排流程，将 LLM 限制在理解、解释和辅助决策，关键动作由规则、权限和状态机约束。",
        "安全与质量：采用 governance → preview/dry-run → execute → verify 链路，配合四级权限、幂等、回读校验、Langfuse、30 条 Golden Cases、LLM-as-judge 和 Playwright E2E。",
    ], "Python、LangGraph、SQL、Qdrant、BGE/fastembed、Langfuse、Streamlit。")
    add_project(doc, "面料图像识别与检索系统", "浙江思维特数字科技有限公司", "2025.10-2026.04", "项目负责人｜带领 2-3 人小组", "为面料研发与选样提供从图片采集、特征提取、向量检索到属性筛选的多模态检索能力。", [
        "数据规模：接入 1,522 个面料对象、4,236 张图片，处理约 14GB 图片数据；组合百度图像检索 API 与本地向量能力。",
        "检索策略：将图像向量相似度与成分、组织、颜色、纹理、克重等业务属性组合，支持结果解释与人工复核。",
        "模型与性能：在 GPU 环境验证 Qwen3-VL-Embedding-8B，使用真实 badcase 对比选型；优化 N+1、图片尺寸和重复计算，加入 Redis 缓存、批量查询与重试。",
        "工程治理：补充 RBAC、数据范围、Celery 定时任务和异步索引，使模型能力具备可运营、可追踪的后台支撑。",
    ], "FastAPI、Python、SQLModel、PostgreSQL、Redis、Celery、百度图像检索 API、Qwen3-VL。")
    add_project(doc, "BugleCat（原 nanocodex 原型）｜Windows 本地 Coding Agent", "个人开源项目", "持续迭代", "独立设计与开发", "验证 Codex 风格 Agent 在本地开发场景中的会话、工具、安全和恢复机制。", [
        "从 Python 原型重构为 Rust Core/CLI + Tauri/Svelte，支持 OpenAI-Compatible、多轮 Agent Loop、流式取消、Tool Calling、MCP、Skills、Resume/Fork 和持久化记忆。",
        "建立 ToolContext/ToolRegistry、Sandbox + Approval、预算与超时、上下文压缩、JSONL 审计、Checkpoint/Restore、隔离 Worker 和 A/B Worktree。",
        "第三方插件不直接进入主进程，按 provider、tools、core、CLI 划分边界，验证企业 Agent 的副作用控制和可回放调试路径。",
    ], "Rust、Python、Tauri、Svelte、MCP、Sandbox、Approval、Worktree。")

    add_section(doc, "To B 平台与架构项目")
    add_project(doc, "Nova CTMS｜CRO 临床试验项目管理系统", "新风新研公司项目", "2026.08", "核心设计与开发", "围绕合同、工时、成本、回款和毛利建立经营数据台账、审批与分析链路。", [
        "拆解报价单解析、列映射、主数据导入、附件预上传及合同/服务/工时/支出等模块，形成可落地的数据模型与接口任务。",
        "实现菜单、页面、Server Action 三层权限守卫，结合项目/中心级数据范围统一审批判定，降低越权风险。",
        "以 Schema 为单一事实源，配合迁移、类型和定向回归校验保障规则一致。",
    ], "Next.js、React、TypeScript、Prisma、Ant Design、XLSX、腾讯云 COS。", compact=True)
    add_project(doc, "Noval Auth｜统一身份与细粒度授权服务", "新风新研公司项目", "2026.08", "架构设计与核心开发", "为 Web、SPA、M2M 和 CTMS 提供统一 OIDC/JWT 身份接入、权限和资源范围控制。", [
        "以 issuer + subject 建立身份映射，通过“身份 → 有效角色 → 权限点 → 资源范围”确定性 SQL 路径完成授权判定。",
        "落地 Logto OIDC/JWT 回调、JWKS 验签、权限同步和授权审计；启动自检、刷新失败撤销和管理写操作恢复采用 Fail-closed 策略。",
    ], "FastAPI、SQLAlchemy、Alembic、PostgreSQL/SQLite、Logto、PyJWT。", compact=True)
    page_break(doc)
    add_section(doc, "To B 平台与架构项目（续）")
    add_project(doc, "标准版 SRM｜Java 21 技术升级", "浙江思维特数字科技有限公司", "2026.01-2026.04", "方案负责人 / 核心开发", "在供应链系统中推进 Java 21 升级、查询链路优化和容器化交付。", [
        "使用 Virtual Threads、Record、Pattern Matching 重构部分查询与 DTO，保持接口契约和业务规则稳定。",
        "通过索引、分页与 SQL 优化改善数万至百万级数据查询；配合 Docker、Jenkins 完成构建、部署与回归。",
    ], "Java 21、Spring Boot、MyBatis、MySQL、Redis、Docker、Jenkins。", compact=True)
    add_project(doc, "金田 SRM｜供应商协同与 ERP 集成", "浙江思维特数字科技有限公司", "2024.09-2026.01", "方案负责人 / 核心开发", "负责供应商协同、金蝶 ERP 双向同步和核心接口性能治理。", [
        "基于 Spring Boot + MyBatis，结合多线程与 Spring 事件监听处理双向同步；通过幂等、重试、校验和异常记录使数据准确率达到 99.9%。",
        "通过索引、SQL 与查询链路优化，将核心接口从 3 秒优化至 500ms。",
    ], "Java、Spring Boot、MyBatis、MySQL、Redis、金蝶 ERP。", compact=True)
    add_project(doc, "集团 CRM｜客户与业务协同平台", "浙江思维特数字科技有限公司", "2023.07-2024.09", "方案负责人 / 核心开发", "建设客户、合同、组织和业务协同模块，统一 ERP、财务等第三方接口规范。", [
        "基于 Spring Boot、Spring Cloud、MyBatis 设计微服务边界，沉淀集成、鉴权、异常处理和日志规则。",
        "参与需求沟通、数据模型、接口设计、任务拆解、联调测试和上线验收。",
    ], "Java、Spring Boot、Spring Cloud、MyBatis、MySQL、Redis。", compact=True)
    add_project(doc, "银联支付微服务与材料商城 B2B 电商", "工路（杭州）信息技术有限公司", "2021.06-2023.07", "Java 研发工程师 / 技术经理", "支撑下单、电子签约、支付和履约的完整 B2B 交易链路。", [
        "以分布式锁、幂等和 RabbitMQ 异步解耦保障支付一致性，支付成功率 99.95%。",
        "集成上上签电子合同与微信支付，支撑日均订单 1000+；负责需求拆解、核心开发、联调、部署和生产排障。",
    ], "Java、Spring Boot、MyBatis、Redis、RabbitMQ、MySQL、Docker、Jenkins。", compact=True)
    add_project(doc, "工业物联网平台｜设备数据采集与可视化", "工路（杭州）信息技术有限公司", "2021-2023", "架构师 / 核心开发", "建设设备接入、数据清洗、时序存储和运营看板链路。", [
        "设计 MQTT / CoAP 接入、Kafka 缓冲、InfluxDB 时序存储和看板查询，梳理设备协议、采集链路、存储模型与展示接口边界。",
    ], "Java、Spring Boot、MQTT、CoAP、Kafka、InfluxDB、Docker。", compact=True)

    add_section(doc, "早期工作项目")
    add_project(doc, "金融交易、银行与政务 OA", "杭州亿凯软件有限公司等", "2015.07-2021.06", "Java 研发 / 模块负责人", "在金融交易、银行智能投顾与政府 OA 项目中积累复杂业务、批处理和现场交付经验。", [
        "参与私募交易系统和金融交易中心骨干项目，完成 Java 后端模块开发与联调交付。",
        "使用 Oracle SQL*Loader 处理批量导入、分片、校验和失败重试，将查询从 10 秒级优化到 1 秒级。",
        "参与农商银行手机银行智能投顾（风险问卷、基金净值、投资建议）及浙江省政府 OA（Word 模板、电子签章、部署支持）。",
    ], "Java、Spring、Oracle、SQL*Loader、MySQL、Linux。")

    page_break(doc)
    add_section(doc, "个人开源与技术预研")
    add_project(doc, "campaign-muti-agent｜多 Agent 协作与治理框架", "个人开源｜github.com/dgy-github/campaign-muti-agent", "持续迭代", "独立开发", "探索复杂任务中的角色协作、事件记录和成本治理。", [
        "以 Coordinator、Executor、Retriever、Reviewer 分工，通过 ROI 路由选择单 Agent 或多 Agent 路径。",
        "使用 append-only EventLog、Policy-as-Code、预算、熔断降级和人工接管，支持回放、恢复与审计。",
    ], "Python、EventLog、Policy-as-Code、RAG、Tool Calling。", compact=True)
    add_project(doc, "Guided Development Project", "个人开源｜github.com/dgy-github/guided-development-project", "持续迭代", "独立开发", "把需求讨论、契约设计、并行开发、测试和交付门禁固化为 G0-G7 工作流。", [
        "以 REQ/API/DB/FE/BE/TEST 追踪需求到实现，采用 API Contract First、Frontend Mock First 和 OpenAPI 生成客户端/MSW。",
        "Agent 提交验证证据后才进入交付，配套项目初始化、能力目录、代码结构检查、集成测试和需求冒烟测试。",
    ], "FastAPI、OpenAPI、Vite、TypeScript、MSW。", compact=True)
    add_project(doc, "Lightweight AI Project Workflow", "个人开源｜github.com/dgy-github/lightweight-ai-project-workflow", "持续迭代", "独立开发", "面向普通人的低判断成本项目管理与 Skill 沉淀方法。", [
        "将讨论清楚 → 计划 → 执行 → 测试验收 → 复用沉淀固化为项目卡、执行清单、验收记录和复用候选卡。",
        "按输入、步骤、产物、验收和返工维度评估 Skill 候选，已沉淀 JD 分析、简历定制、证据核验和投递草稿能力。",
    ], "Markdown、Python、Skills、项目治理。", compact=True)
    add_project(doc, "MicrocodeX Short Drama Studio", "个人开源（Alpha）｜github.com/dgy-github/microcodex-short-drama-studio", "持续迭代", "独立开发", "Windows-first 短剧创作工作台，验证固定 DAG、事件回放和版本化 Story Package。", [
        "Rust 负责产品契约、事件协议、Provider/存储边界与评测；Python sidecar 通过 campaign-muti-agent 编排固定 17-task DAG。",
        "支持审批、导出、Fail-closed 和 SSE 可恢复回放；一次真实运行完成 17/17 任务、6 集、154K tokens（Alpha / advisory）。",
    ], "Rust、Python、Tauri、Svelte、DAG、EventLog、SSE。", compact=True)
    add_project(doc, "Story Image Agent / Story Video Agent", "个人开源与预研｜github.com/dgy-github/story-image-agent / story-video-agent", "持续迭代", "独立开发", "以 JSON Schema、Provider Adapter 和质量证据验证可审计的故事生图/生视频工作流。", [
        "固定 source_spans、artifact lineage、stage receipt 与质量门禁；图片工作流支持 story alignment、composition、identity consistency 和 artifact-free 评估。",
        "视频工作流覆盖粗生成、裁剪、补段、评估、精生成、checkpoint、幂等、取消、重试与恢复；默认 simulator/FakeProvider/FFmpeg，不夸大为线上云端规模。",
    ], "Python、JSON Schema、Provider Adapter、FFmpeg、Mock/Simulator。", compact=True)
    add_project(doc, "Windows Computer-Use MCP / 本地模型与 Agent 评测", "个人技术预研", "持续迭代", "独立开发", "验证桌面自动化和本地模型在安全执行、结构化输出与隐私场景中的边界。", [
        "基于 UIA、OCR、分级定位、allow-list、文本扫描、人工 Approval、前后截图和操作日志，构建可回退的桌面操作链路。",
        "使用 Qwen、llama-server、OpenAI-compatible API、Golden Cases 和失败样本比较意图识别、工具调用、结构化输出、成本与延迟。",
    ], "Windows UIA、OCR、MCP、Qwen、llama-server、Golden Cases。", compact=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
