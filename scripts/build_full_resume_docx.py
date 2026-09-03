"""Build the full project-based resume tailored to the supplied AI Agent JD."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(r"D:/github_dgy/output/docx/狄广宇-AI-Agent岗位AI强化版.docx")
BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = RGBColor(117, 117, 117)
FONT = "Microsoft YaHei"


def set_run(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_border(paragraph, color="1F4E79", size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, 15, True, BLUE)
    set_cell_border(p)
    return p


def add_lead(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.keep_together = True
    r = p.add_run(label)
    set_run(r, 10.5, True)
    r = p.add_run(text)
    set_run(r, 10.5)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Cm(0.55 + level * 0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    # When a bullet contains a category prefix (e.g. “架构：” or “评测：”),
    # emphasize that prefix so AI capability and engineering evidence are easy to scan.
    if "：" in text and text.index("：") <= 8:
        prefix, rest = text.split("：", 1)
        r = p.add_run(prefix + "：")
        set_run(r, 10.1, True, BLUE)
        r = p.add_run(rest)
        set_run(r, 10.1)
    else:
        r = p.add_run(text)
        set_run(r, 10.1)
    return p


def add_project(doc, title, meta, summary, bullets, stack=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run(r, 11.5, True, BLUE)
    if meta:
        r = p.add_run("  " + meta)
        set_run(r, 9.2, False, LIGHT_GRAY)
    add_lead(doc, "项目定位：", summary)
    for bullet in bullets:
        add_bullet(doc, bullet)
    if stack:
        add_lead(doc, "技术栈：", stack)


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("狄广宇")
    set_run(r, 23, True, BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Java 后端 / 技术负责人 / AI 应用工程化")
    set_run(r, 14, False, GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("电话 19290569945  ·  邮箱 2537676793@qq.com  ·  GitHub github.com/dgy-github")
    set_run(r, 10, False, GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("12 年研发经验  ·  9 年 Java 后端及架构经验  ·  约 2 年 AI 应用工程化实战  ·  本科")
    set_run(r, 10.5, True)

    add_section_title(doc, "职业概述")
    add_bullet(doc, "12 年研发经验，其中 9 年 Java 后端与架构经验；具备 To B 产品从需求沟通、业务建模、方案设计、开发测试到上线迭代的完整交付经历。")
    add_bullet(doc, "长期负责供应链、SRM、ERP、CRM、B2B 电商、支付和内部管理系统，熟悉复杂业务规则抽象、接口集成、数据同步和系统稳定性建设。")
    add_bullet(doc, "近年扩展 AI 应用工程化能力，完成制造业排产 Agent、面料图像检索、Coding Agent 等项目，将 AI 与已有业务系统和工程流程结合。")
    add_bullet(doc, "带领 2-3 人小组（1 名实习生、1 名前端成员），负责方案评审、任务拆解、代码 Review、进度推进和交付验收；长期使用 Docker、Jenkins。")

    add_section_title(doc, "能力概览")
    add_lead(doc, "后端与架构：", "Java 8/11/21、JVM、多线程、IO/NIO、Spring Boot、Spring Cloud、MyBatis、微服务、分布式事务、幂等、事件驱动。")
    add_lead(doc, "企业系统：", "SRM、ERP、CRM、B2B 电商、支付、供应链、物联网、权限系统；擅长业务拆解、数据建模、接口设计、第三方集成和性能优化。")
    add_lead(doc, "AI 应用：", "LangGraph、RAG、Tool Calling、MCP、Workflow、Prompt 工程、多模态检索、Agent 评测、Langfuse、Golden Cases。")
    add_lead(doc, "交付与管理：", "2-3 人团队管理、代码 Review、任务拆解、Docker、Jenkins、测试验收、日志追踪、权限审计、失败恢复。")

    add_section_title(doc, "技术栈")
    add_lead(doc, "Java 后端：", "Java 8/11/21、Spring Boot、Spring Cloud、MyBatis、Flowable、FastAPI、微服务、JVM、多线程、IO/NIO。")
    add_lead(doc, "数据与基础设施：", "MySQL、Oracle、PostgreSQL、Redis、RabbitMQ、SQL 优化、数据建模、任务调度、Docker、Jenkins。")
    add_lead(doc, "AI 与前端：", "LangGraph、RAG、Tool Calling、MCP、Qdrant、BGE、Qwen、Langfuse、React、Next.js、TypeScript、Ant Design。")

    add_section_title(doc, "工作经历")
    add_lead(doc, "浙江思维特数字科技有限公司｜AI 应用研发 / 技术经理｜2023.07-2026.05", "  负责产品线技术方案、架构规划与 AI 应用落地，带领小组推进制造业 Agent、SRM、ERP、CRM 等系统。")
    add_lead(doc, "工路（杭州）信息技术有限公司｜Java 研发工程师 / 技术经理｜2021.06-2023.07", "  负责支付微服务、物联网平台和 B2B 电商系统的架构与核心开发。")
    add_lead(doc, "杭州亿凯软件有限公司等｜Java 软件工程师｜2015.07-2021.06", "  参与金融、银行、政务系统研发，积累 Java 后端、数据处理和现场交付经验。")

    page_break(doc)
    add_section_title(doc, "核心项目经历")
    add_project(
        doc,
        "Nova CTMS｜CRO 临床试验项目管理系统",
        "新风新研公司项目｜2026.08｜核心设计与开发",
        "以合同、服务清单、项目/日常工时、支出、回款、开票、成本和毛利为业务追溯主链，建设经营数据台账、审批和分析一体化的 To B 管理系统。",
        [
            "将合同、服务事项、项目成员、工时填报、回款开票和实际支出拆解为模块、数据模型与接口任务，支持报价单解析、列映射、主数据导入和附件预上传。",
            "实现菜单、页面、Server Action 三层权限守卫，结合项目/中心级数据范围，统一“待我审批”判定，降低越权和跨范围访问风险。",
            "将服务事项、角色承诺工时、审批后实际工时和合同成本归集到实时毛利分析，成本脱敏逻辑下沉数据层。",
            "以 Schema 作为单一事实源，通过迁移、类型与定向回归校验保障数据和业务规则一致。",
        ],
        "Next.js 15、React 19、TypeScript、Prisma 6、Ant Design 6、XLSX、腾讯云 COS。",
    )

    add_project(
        doc,
        "Noval Auth｜统一身份、应用接入与细粒度授权服务",
        "新风新研公司项目｜2026.08｜架构设计与核心开发",
        "将身份提供方与业务授权解耦，为 Web、SPA、M2M 和 CTMS 提供统一登录、角色权限和租户/应用/项目/中心数据范围。",
        [
            "以 issuer + subject 建立身份映射，通过“身份 → 有效角色 → 权限点 → 资源范围”的确定性 SQL 路径完成授权判定。",
            "完成 Logto OIDC/JWT 登录回调、JWKS 验签、权限导入同步、管理操作及授权决策审计。",
            "采用 Fail-closed 启动自检；开发身份仅允许本机；会话只保存签名编号，刷新失败立即撤销，管理写操作与持久化状态支持失败恢复。",
        ],
        "FastAPI、SQLAlchemy、Alembic、PostgreSQL/SQLite、Logto OIDC/JWT、PyJWT。",
    )

    add_project(
        doc,
        "标准版 SRM｜Java 21 技术升级",
        "浙江思维特数字科技有限公司｜2026.01-2026.04｜方案负责人 / 核心开发",
        "在供应链系统中推进 Java 21 技术升级和查询链路优化，验证新版本语言特性在企业级后台系统中的落地方式。",
        [
            "使用 Virtual Threads、Record、Pattern Matching 重构部分查询和数据传输代码，保持接口契约和业务规则稳定。",
            "针对核心查询链路进行索引、分页和 SQL 优化，支持数万到百万级数据查询场景。",
            "配合 Docker、Jenkins 完成构建、部署和回归验证；环境配置、部署脚本和流水线维护过程中使用 AI 辅助，并保留人工检查与回滚边界。",
        ],
        "Java 21、Spring Boot、MyBatis、MySQL、Redis、Docker、Jenkins。",
    )

    add_project(
        doc,
        "面料图像识别与检索系统",
        "浙江思维特数字科技有限公司｜2025.10-2026.04｜项目负责人，带领 2-3 人小组",
        "面向面料研发和业务选样，建设从图片采集、特征提取、向量检索到结果筛选的多模态图像检索系统。",
        [
            "数据规模：接入 1,522 个面料对象、4,236 张图片，处理约 14GB 图片数据；基于百度图像检索 API 和本地向量库实现以图搜图及相似面料检索。",
            "多模态检索：围绕成分、组织、颜色、纹理、克重等字段建立结构化筛选条件，将图像向量相似度与业务属性过滤组合，支持检索结果解释和人工复核。",
            "模型验证：在 GPU 环境验证 Qwen3-VL-Embedding-8B 等多模态模型，通过真实 badcase 对比模型选型，并记录 token 失效、召回偏差等问题。",
            "性能优化：针对检索接口中的 N+1 查询、图片尺寸过大和重复计算进行优化，增加图片裁剪压缩、Redis 缓存、批量查询和失败重试。",
            "工程治理：补充 RBAC、数据范围控制和 Celery 定时任务，保证批量索引、异步处理和后台管理的可追溯性。",
        ],
        "FastAPI、Python、SQLModel、PostgreSQL、Redis、Celery、百度图像检索 API、Qwen3-VL、向量检索。",
    )

    add_project(
        doc,
        "制造业排产 Agent",
        "浙江思维特数字科技有限公司｜2025-2026.04｜技术负责人，带领 2-3 人小组",
        "在已有生产排产系统上完成 Agent 化重构，交付自然语言订单查询、机台负载分析、规格/颜色机台推荐、排产诊断和安全写操作。",
        [
            "数据与路由：基于 49,981 条真实生产订单抽象 20 类意图路由，以规则和 SQL 优先处理确定性请求，使约 80% 请求无需调用 LLM，离线测试意图识别准确率约 95%。",
            "业务建模：围绕订单、机台、规格、颜色和交期建立业务上下文，覆盖 75D / 100D / 300D 等规格排产，以及机台负载、替代机台和排产原因解释。",
            "Agent 编排：使用 LangGraph StateGraph 编排固定业务流程，将 LLM 限制在意图理解、解释和辅助决策环节，关键业务动作由规则、权限和状态机约束。",
            "检索策略：按查询类型组合 SQL 精确查询、Hybrid 检索、Composite 多源融合和 GraphRAG，支持订单、机台、规格及替代机台等复杂业务问答。",
            "安全执行：建立 governance → preview（dry-run）→ execute → verify 写操作安全链路，加入四级权限、幂等、一致性及回读校验，避免模型直接控制生产流程。",
            "评测与观测：通过 Langfuse、30 条 Golden Cases、LLM-as-judge 与 Playwright E2E 建立可观测和回归体系；已部署内网并完成核心链路验证。",
            "交付状态：2026.04 前完成查询与诊断链路验证，写操作与真实排产 API 解耦，按阶段推进灰度接入。",
        ],
        "LangGraph、Python 3.12、SQL、SQLite、Qdrant、BGE/fastembed、Langfuse、Streamlit。",
    )

    add_project(
        doc,
        "金田 SRM｜供应商协同与 ERP 集成",
        "浙江思维特数字科技有限公司｜2024.09-2026.01｜方案负责人 / 核心开发",
        "面向供应商协同和采购业务，负责 SRM 后端模块、金蝶 ERP 双向同步及核心接口性能优化。",
        [
            "基于 Spring Boot + MyBatis 与金蝶 ERP 进行双向数据同步，使用多线程和 Spring 事件监听处理数据，数据准确率达到 99.9%。",
            "围绕同步幂等、失败重试、数据校验和异常记录设计处理链路，保证主数据和业务单据可追溯。",
            "通过索引、SQL 和查询链路优化，将核心接口响应从 3 秒优化到 500ms。",
        ],
        "Java、Spring Boot、MyBatis、MySQL、Redis、金蝶 ERP 集成。",
    )

    add_project(
        doc,
        "集团 CRM｜客户与业务协同平台",
        "浙江思维特数字科技有限公司｜2023.07-2024.09｜方案负责人 / 核心开发",
        "面向客户、合同和业务协同场景，负责 CRM 微服务架构、接口规范和核心业务模块设计。",
        [
            "基于 Spring Boot、Spring Cloud、MyBatis 设计微服务架构，拆分客户、合同、组织和业务协同模块。",
            "统一 ERP、财务等第三方接口规范，沉淀可复用的集成、鉴权、异常处理和日志记录规则。",
            "参与需求沟通、数据模型与接口设计、任务拆解、联调测试和上线验收。",
        ],
        "Java、Spring Boot、Spring Cloud、MyBatis、MySQL、Redis、ERP / 财务系统集成。",
    )

    add_project(
        doc,
        "银联支付微服务与材料商城 B2B 电商",
        "工路（杭州）信息技术有限公司｜2021.06-2023.07｜Java 研发工程师 / 技术经理",
        "负责支付微服务、订单合同和 B2B 电商后台的架构与核心开发，支撑从下单、签约到支付的完整链路。",
        [
            "支付微服务采用分布式锁、幂等和 RabbitMQ 异步解耦保障支付一致性，支付成功率 99.95%。",
            "材料商城 B2B 电商基于 Spring Boot + MyBatis，集成上上签电子合同和微信支付，支撑订单、合同、支付流程，日均订单 1000+。",
            "负责需求拆解、接口设计、核心代码开发、联调测试、部署上线和生产问题排查。",
        ],
        "Java、Spring Boot、MyBatis、Redis、RabbitMQ、MySQL、微信支付、上上签、Docker、Jenkins。",
    )

    page_break(doc)
    add_section_title(doc, "早期工作项目（按时间倒序）")
    add_project(
        doc,
        "工业物联网平台｜设备数据采集与可视化",
        "工路（杭州）信息技术有限公司｜2021-2023｜架构师 / 核心开发",
        "面向设备联网、生产数据采集和运营看板，负责平台整体架构、数据链路和关键模块开发。",
        [
            "设计 MQTT / CoAP 接入、Kafka 消息缓冲、InfluxDB 时序存储和看板查询链路，完成设备数据采集、清洗、存储与可视化的一体化落地。",
            "参与项目预研与架构方案设计，梳理设备协议、采集链路、存储模型和展示接口之间的边界。",
            "负责关键模块实现、接口联调、上线排障和现场技术支持。",
        ],
        "Java、Spring Boot、MQTT、CoAP、Kafka、InfluxDB、MySQL、Redis、Docker。",
    )

    add_project(
        doc,
        "金融交易与银行业务系统",
        "杭州亿凯软件有限公司等｜2015-2021｜Java 研发 / 模块负责人",
        "参与私募交易、银行和政务系统建设，承担后端模块、批量数据处理、查询性能和现场交付工作。",
        [
            "参与私募交易系统以及金融交易中心骨干项目，完成相关 Java 后端模块开发和联调交付。",
            "使用 Oracle SQL*Loader 优化批量数据导入，处理大批量数据分片、校验和失败重试；通过索引、分页和 SQL 改写将查询从 10 秒级降至 1 秒级。",
            "参与农商银行手机银行智能投顾模块，覆盖风险问卷、基金净值与投资建议等业务功能。",
            "参与浙江省政府 OA 公文系统，完成 Word 模板和电子签章相关功能，并承担现场部署与技术支持。",
        ],
        "Java、Spring、Oracle、SQL*Loader、MySQL、REST API、Linux。",
    )

    page_break(doc)
    add_section_title(doc, "个人项目 / 开源项目")

    add_project(
        doc,
        "BugleCat｜Windows 本地 Coding Agent",
        "个人开源项目｜Python / Rust / Tauri + Svelte",
        "从零实现 Codex 风格多轮 Agent Loop，并从 Python 原型重构为 Rust Core/CLI 与 Tauri 桌面端，用于验证企业级 Agent 的工具边界、会话恢复和安全执行。",
        [
            "实现 OpenAI-Compatible 模型、流式输出/取消、工具调用、MCP、Skills、持久化记忆、会话 Resume/Fork、定时任务和 A/B Worktree 对比。",
            "按 sandbox、provider、tools、core、CLI 划分职责；模型、工具、会话和记忆通过显式类型契约解耦，第三方插件不直接进入主进程。",
            "建立 Sandbox + Approval 状态机、调用/工具/上下文预算、超时控制、JSONL 日志、会话快照、每轮 Checkpoint、安全恢复及隔离 Worktree 回归，控制副作用和状态漂移。",
        ],
        "Rust、Python、Tauri、Svelte、MCP、Agent Loop、Sandbox、Approval、Worktree。",
    )

    add_project(
        doc,
        "campaign-muti-agent｜多 Agent 协作与治理框架",
        "个人开源 / 技术预研｜github.com/dgy-github/campaign-muti-agent",
        "探索复杂任务中多角色 Agent 的协作、路由、事件记录和成本治理，沉淀可迁移到企业 AI 应用的基础模式。",
        [
            "以 Coordinator、Executor、Retriever、Reviewer 进行角色分工，通过 ROI 路由选择单 Agent 或多 Agent 路径。",
            "使用 append-only EventLog 记录任务事件，支持状态回放、失败恢复和过程审计，避免只保留最终文本导致无法定位问题。",
            "引入 Policy-as-Code、预算控制、熔断降级和人工接管，约束工具权限、调用次数和高风险动作。",
        ],
        "Python、Agent Loop、EventLog、Policy-as-Code、RAG、Tool Calling。",
    )

    add_project(
        doc,
        "Windows Computer-Use MCP｜桌面操作安全执行",
        "个人技术预研",
        "验证 Windows 桌面 UI 自动化和 Computer-Use MCP 在真实操作中的定位、确认与回滚边界。",
        [
            "结合 UIA、OCR、分级定位和多次确认处理窗口、控件及文本目标，降低单次坐标定位带来的误操作风险。",
            "对点击、输入、文件和提交等动作设置 allow-list、文本扫描和人工 Approval；未确认或定位不确定时保持不执行。",
            "补充前后状态截图、操作日志和失败回退路径，为桌面 Agent 的可追溯验收提供证据。",
        ],
        "Windows UIA、OCR、MCP、Computer-Use、allow-list、Approval。",
    )

    add_project(
        doc,
        "本地模型与 Agent 评测预研",
        "个人技术预研｜Qwen / llama-server",
        "验证本地模型、OpenAI 兼容接口与 Agent 工作流的组合方式，为内网部署和数据敏感场景提供选型依据。",
        [
            "使用 Qwen 系列模型和 llama-server 提供 OpenAI 兼容接口，比较本地模型在意图识别、工具调用和结构化输出中的可用性。",
            "通过 Golden Cases、失败样本和人工复核记录模型差异，区分模型能力问题、提示词问题和工具链问题。",
            "围绕成本、延迟、隐私、可观测性和故障恢复整理本地化 Agent 路径，作为企业 AI 应用落地的预研依据。",
        ],
        "Qwen、llama-server、OpenAI-compatible API、Prompt Engineering、Golden Cases。",
    )

    add_section_title(doc, "教育背景")
    add_lead(doc, "苏州工学院｜软件工程｜本科（统招）｜2010.09-2014.06", "")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("狄广宇｜AI Agent 岗位定制简历")
    set_run(r, 8, False, LIGHT_GRAY)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
