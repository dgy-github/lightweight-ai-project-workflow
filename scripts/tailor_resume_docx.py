"""Create a tailored copy of the supplied resume with confirmed facts only."""

from pathlib import Path
import shutil

from docx import Document


SOURCE = Path(r"D:/github_dgy/tmp/resume-source.docx")
OUTPUT = Path(r"D:/github_dgy/output/docx/狄广宇-AI-Agent岗位定制版.docx")


def replace_all(document: Document, old: str, new: str) -> int:
    count = 0
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
    return count


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)
    replacements = {
        "11 年企业级后端工程经验": "12 年研发经验，9 年 Java 后端及架构经验",
        "11 年后端工程能力迁移到 AI": "9 年 Java 后端工程能力迁移到 AI",
        "11 年企业级系统的工程化能力": "9 年 Java 后端与企业级系统工程化能力",
        "Spring Boot / Cloud / Flowable": "Spring Boot / Spring Cloud / MyBatis / Flowable",
        "Docker；第三方集成": "Docker / Jenkins（长期使用，借助 AI 完成环境配置、部署和流水线维护）；第三方集成",
        "与金蝶 ERP 双向数据同步": "基于 Spring Boot + MyBatis，与金蝶 ERP 双向数据同步",
        "主导 Spring Cloud 微服务架构": "基于 Spring Boot / Spring Cloud / MyBatis 主导微服务架构",
        "分布式锁 + 幂等性方案保障支付一致性": "基于 Spring Boot + MyBatis，采用分布式锁 + 幂等性方案保障支付一致性",
        "集成上上签电子合同、微信支付": "基于 Spring Boot + MyBatis，集成上上签电子合同、微信支付",
    }
    for old, new in replacements.items():
        replace_all(document, old, new)

    for paragraph in document.paragraphs:
        if "AI 应用工程师 / Agent Engineer" in paragraph.text:
            paragraph.text = "AI 应用工程师 / Agent Engineer"
        if paragraph.text.startswith("苏州思为特数字科技有限公司"):
            paragraph.text = paragraph.text.replace("2023.07 - 至今", "2023.07 - 至今")
        if paragraph.text.startswith("浙江思为特数字科技有限公司"):
            paragraph.text = paragraph.text.replace("AI 应用研发 / 技术经理", "AI 应用研发 / 技术经理（带领 2-3 人 AI 小组）")
        if paragraph.text.startswith("AI 团队建设") and "2-3 人" not in paragraph.text:
            for run in paragraph.runs:
                if run.text.startswith("2024 年起"):
                    run.text = "带领 2-3 人小组（1 名实习生、1 名前端成员）；" + run.text
        if "核心项目：制造业排产 Agent" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace("技术负责人，带 1 名实习生", "技术负责人，带领 2-3 人小组（1 名实习生、1 名前端成员）")
        if "制造业排产 Agent" in paragraph.text and "浙江思为特" not in paragraph.text:
            paragraph.add_run("；项目归属：浙江思为特数字科技有限公司")
        if "Nova CTMS" in paragraph.text and "新风新研" not in paragraph.text:
            paragraph.add_run("；项目归属：新风新研公司项目")

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
